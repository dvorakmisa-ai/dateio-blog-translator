import os
import re
import json
import subprocess
from datetime import datetime
from urllib.parse import urljoin

import requests
from bs4 import BeautifulSoup
from openai import OpenAI
from docx import Document


BLOG_INDEX = "https://www.dateioplatform.com/resources/blog"
STATE_FILE = "state.json"

# Kam ukládat Word soubory v repozitáři:
OUTPUT_DIR = "translations/HU"


# ----------------------
# Helpers
# ----------------------
def require_env(name: str) -> str:
    value = os.getenv(name)
    if not value:
        raise RuntimeError(
            f"Chybí proměnná prostředí {name}. "
            f"V GitHub Actions ji předej přes env: {name}: ${{{{ secrets.{name} }}}}"
        )
    return value


def fetch_html(url: str) -> str:
    r = requests.get(url, timeout=30, headers={"User-Agent": "Mozilla/5.0"})
    r.raise_for_status()
    return r.text


def load_state() -> set[str]:
    if not os.path.exists(STATE_FILE):
        return set()
    with open(STATE_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
    return set(data.get("processed_urls", []))


def save_state(processed_urls: set[str]) -> None:
    with open(STATE_FILE, "w", encoding="utf-8") as f:
        json.dump({"processed_urls": sorted(processed_urls)}, f, ensure_ascii=False, indent=2)


def extract_post_urls_from_index(index_html: str) -> list[str]:
    soup = BeautifulSoup(index_html, "html.parser")
    urls = set()
    for a in soup.select("a[href]"):
        href = a.get("href", "")
        if "/resources/post/" in href:
            urls.add(urljoin(BLOG_INDEX, href))
    return sorted(urls)


def extract_article(article_html: str, article_url: str) -> dict:
    soup = BeautifulSoup(article_html, "html.parser")

    title = "Untitled"
    h1 = soup.find("h1")
    if h1 and h1.get_text(strip=True):
        title = h1.get_text(strip=True)
    elif soup.title and soup.title.get_text(strip=True):
        title = soup.title.get_text(strip=True)

    # meta title: og:title -> <title>
    meta_title = None
    og_title = soup.find("meta", attrs={"property": "og:title"})
    if og_title and og_title.get("content"):
        meta_title = og_title["content"].strip()
    elif soup.title and soup.title.get_text(strip=True):
        meta_title = soup.title.get_text(strip=True)

    # meta description: description -> og:description
    meta_description = None
    meta_desc = soup.find("meta", attrs={"name": "description"})
    if meta_desc and meta_desc.get("content"):
        meta_description = meta_desc["content"].strip()
    else:
        og_desc = soup.find("meta", attrs={"property": "og:description"})
        if og_desc and og_desc.get("content"):
            meta_description = og_desc["content"].strip()

    paragraphs = [
        p.get_text(" ", strip=True)
        for p in soup.find_all("p")
        if p.get_text(strip=True)
    ]
    body = "\n\n".join(paragraphs).strip()

    return {
        "url": article_url,
        "title": title,
        "meta_title": meta_title,
        "meta_description": meta_description,
        "body": body,
    }


def translate_to_hungarian(client: OpenAI, article: dict) -> dict:
    prompt = f"""
Přelož následující obsah do maďarštiny.

Vrať výstup POUZE jako JSON:
{{
  "title": "...",
  "meta_title": "...",
  "meta_description": "...",
  "body": "..."   // Markdown
}}

Pravidla:
- zachovej význam a marketingový tón
- meta_title ideálně do 60 znaků
- meta_description ideálně do 160 znaků
- body vrať v Markdownu, zachovej odstavce
- pokud meta_title nebo meta_description chybí, navrhni je
- žádné jiné klíče, žádné komentáře

# TITLE
{article["title"]}

# META TITLE
{article.get("meta_title") or ""}

# META DESCRIPTION
{article.get("meta_description") or ""}

# BODY
{article["body"]}
""".strip()

    resp = client.responses.create(
        model=os.getenv("OPENAI_MODEL", "gpt-4.1-mini"),
        input=prompt,
    )
    text = resp.output_text.strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", text, flags=re.DOTALL)
        if not m:
            raise RuntimeError(f"Model nevrátil validní JSON. Výstup:\n{text}")
        return json.loads(m.group(0))


def markdownish_to_docx_paragraphs(doc: Document, body_md: str) -> None:
    # Jednoduchý převod: rozdělení po odstavcích, odstranění # headingů
    body_md = (body_md or "").strip()
    if not body_md:
        doc.add_paragraph("(empty)")
        return

    for block in body_md.split("\n\n"):
        t = block.strip()
        if not t:
            continue
        t = re.sub(r"^#{1,6}\s+", "", t)  # remove markdown headings
        doc.add_paragraph(t)


def build_docx_file(original: dict, hu: dict, out_path: str) -> None:
    doc = Document()
    doc.add_heading((hu.get("title") or original["title"]).strip(), level=1)

    doc.add_paragraph(f"Originál: {original['url']}")
    doc.add_paragraph("")

    doc.add_heading("SEO meta (HU)", level=2)
    doc.add_paragraph(f"Meta title: {(hu.get('meta_title') or '').strip()}")
    doc.add_paragraph(f"Meta description: {(hu.get('meta_description') or '').strip()}")

    doc.add_paragraph("")
    doc.add_heading("Obsah (HU)", level=2)
    markdownish_to_docx_paragraphs(doc, hu.get("body") or "")

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)


def run(cmd: list[str]) -> None:
    subprocess.run(cmd, check=True)


def git_commit_and_push(file_paths: list[str], message: str) -> bool:
    """
    Vrátí True, pokud se něco commitlo. False pokud nebyly změny.
    """
    # Nastavení identity
    run(["git", "config", "user.email", "github-actions[bot]@users.noreply.github.com"])
    run(["git", "config", "user.name", "github-actions[bot]"])

    # Add files
    run(["git", "add", *file_paths])

    # Pokud není co commitnout, git commit skončí s code 1 – ošetříme
    res = subprocess.run(["git", "commit", "-m", message], capture_output=True, text=True)
    if res.returncode != 0:
        # typicky: "nothing to commit"
        out = (res.stdout or "") + "\n" + (res.stderr or "")
        if "nothing to commit" in out.lower():
            return False
        raise RuntimeError(f"git commit failed:\n{out}")

    # Push (použije credentials z actions/checkout)
    run(["git", "push"])
    return True


def github_blob_url(repo: str, branch: str, path: str) -> str:
    # repo = "owner/name"
    return f"https://github.com/{repo}/blob/{branch}/{path}"


def send_teams_message(title_hu: str, meta_title_hu: str, meta_desc_hu: str, original_url: str, file_url: str) -> None:
    webhook = require_env("TEAMS_WEBHOOK_URL").strip()
    payload = {
        "text": (
            "📄 **HU překlad blogu – Word soubor**\n\n"
            f"**Název (HU):** {title_hu}\n\n"
            f"**Meta title (HU):** {meta_title_hu}\n\n"
            f"**Meta description (HU):** {meta_desc_hu}\n\n"
            f"🔗 **Originál:** {original_url}\n\n"
            f"📎 **DOCX v repu:** {file_url}\n"
        )
    }
    r = requests.post(webhook, json=payload, headers={"Content-Type": "application/json"}, timeout=30)
    if r.status_code >= 400:
        print("TEAMS ERROR:", r.status_code, r.text)
    r.raise_for_status()


def main():
    require_env("OPENAI_API_KEY")
    require_env("TEAMS_WEBHOOK_URL")

    # GitHub kontext (Actions to nastavuje automaticky)
    repo = require_env("GITHUB_REPOSITORY")        # "owner/repo"
    branch = os.getenv("GITHUB_REF_NAME", "main")  # často "main"

    processed = load_state()

    index_html = fetch_html(BLOG_INDEX)
    post_urls = extract_post_urls_from_index(index_html)
    new_urls = [u for u in post_urls if u not in processed]

    if not new_urls:
        print("Žádné nové články.")
        save_state(processed)  # vždy vytvoří state.json
        return

    client = OpenAI(api_key=require_env("OPENAI_API_KEY"))

    for url in new_urls:
        article_html = fetch_html(url)
        article = extract_article(article_html, url)

        if not article["body"]:
            print(f"Nelze extrahovat text: {url}")
            processed.add(url)
            continue

        hu = translate_to_hungarian(client, article)

        # Sestavení názvu souboru
        slug = re.sub(r"[^a-zA-Z0-9_-]+", "-", url.split("/")[-1]).strip("-")
        dt = datetime.now().strftime("%Y-%m-%d")
        filename = f"{dt}-{slug}-HU.docx"

        rel_path = f"{OUTPUT_DIR}/{filename}"
        out_path = os.path.join(os.getcwd(), rel_path)

        build_docx_file(article, hu, out_path)

        # commit + push docx + state
        save_state(processed | {url})
        committed = git_commit_and_push(
            [rel_path, STATE_FILE],
            message=f"Add HU translation DOCX: {filename}",
        )

        file_url = github_blob_url(repo, branch, rel_path)

        send_teams_message(
            title_hu=(hu.get("title") or article["title"]).strip(),
            meta_title_hu=(hu.get("meta_title") or "").strip(),
            meta_desc_hu=(hu.get("meta_description") or "").strip(),
            original_url=url,
            file_url=file_url,
        )

        print(f"Hotovo: {rel_path} (committed={committed})")
        processed.add(url)

    save_state(processed)


if __name__ == "__main__":
    main()
